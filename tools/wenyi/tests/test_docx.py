"""DOCX 输入解析、组装与 CLI 默认导出格式。"""

from __future__ import annotations

import os
import tempfile
import unittest
from unittest.mock import patch

from docx import Document as DocxDocument
from docx.shared import RGBColor
from typer.testing import CliRunner

from trans_novel.assemble.docx_writer import _assemble_docx
from trans_novel.assemble.writer import assemble
from trans_novel.cli import _resolve_output_format, app
from trans_novel.config import Config
from trans_novel.ingest.docx_reader import read_docx
from trans_novel.ingest.models import KIND_HEADING, KIND_TEXT
from trans_novel.ingest.segmenter import load_document
from trans_novel.pipeline.docx_styles import (
    merge_align_results,
    proportional_range_placements,
)
from trans_novel.pipeline.runstore import STATUS_DONE, RunStore


def _write_sample_docx(path: str) -> None:
    doc = DocxDocument()
    doc.add_heading("Chapter One", level=1)
    doc.add_paragraph("Hello world.")
    doc.add_heading("Section", level=2)
    doc.add_paragraph("More text.")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "A"
    table.cell(0, 1).text = "B"
    table.cell(1, 0).text = "C"
    table.cell(1, 1).text = "D"
    doc.save(path)


class TestDocxReader(unittest.TestCase):
    def test_read_headings_paragraphs_and_table(self):
        with tempfile.TemporaryDirectory() as directory:
            path = os.path.join(directory, "sample.docx")
            _write_sample_docx(path)
            book = read_docx(path, "en", "zh")
        self.assertEqual(book.fmt, "docx")
        self.assertEqual(len(book.chapters), 1)
        self.assertEqual(book.chapters[0].title, "Chapter One")
        kinds = [s.kind for s in book.chapters[0].segments]
        self.assertEqual(kinds[0], KIND_HEADING)
        self.assertEqual(book.chapters[0].segments[0].meta.get("heading_level"), 1)
        self.assertEqual(kinds[1], KIND_TEXT)
        table_segs = [s for s in book.chapters[0].segments if s.meta.get("table_id") == 0]
        self.assertEqual(len(table_segs), 4)
        self.assertEqual(
            {(s.meta["row"], s.meta["col"]): s.source for s in table_segs}[(0, 0)], "A"
        )

    def test_load_document_routes_docx(self):
        with tempfile.TemporaryDirectory() as directory:
            path = os.path.join(directory, "sample.docx")
            _write_sample_docx(path)
            book = load_document(path, "en", "zh")
        self.assertEqual(book.fmt, "docx")

    def test_uniform_style_skips_mixed_meta(self):
        with tempfile.TemporaryDirectory() as directory:
            path = os.path.join(directory, "styled.docx")
            doc = DocxDocument()
            paragraph = doc.add_paragraph()
            run = paragraph.add_run("All bold blue")
            run.bold = True
            run.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)
            doc.save(path)
            book = read_docx(path, "en", "zh")
        segment = book.chapters[0].segments[0]
        self.assertEqual(segment.meta.get("docx_style", {}).get("bold"), True)
        self.assertEqual(segment.meta.get("docx_style", {}).get("color"), "0000FF")
        self.assertNotIn("docx_styles", segment.meta)

    def test_mixed_runs_record_style_spans(self):
        with tempfile.TemporaryDirectory() as directory:
            path = os.path.join(directory, "mixed.docx")
            doc = DocxDocument()
            paragraph = doc.add_paragraph()
            red = paragraph.add_run("Red ")
            red.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
            paragraph.add_run("plain")
            doc.save(path)
            book = read_docx(path, "en", "zh")
        segment = book.chapters[0].segments[0]
        self.assertNotIn("docx_style", segment.meta)
        items = segment.meta["docx_styles"]["items"]
        self.assertEqual(len(items), 1)
        self.assertEqual(items[0]["source_start"], 0)
        self.assertEqual(items[0]["source_end"], 4)
        self.assertEqual(items[0]["color"], "FF0000")

    def test_toc_line_with_visible_number_skips_list_meta(self):
        """目录「1. Title」正文已含序号时，不应再套自动编号。"""
        with tempfile.TemporaryDirectory() as directory:
            path = os.path.join(directory, "toc.docx")
            doc = DocxDocument()
            doc.add_heading("Contents", level=1)
            # 模拟带失效 numPr 的目录行：正文已有「1.」
            doc.add_paragraph("1. The myth of primitive society", style="List Number")
            doc.save(path)
            book = read_docx(path, "en", "zh")
            toc = next(s for s in book.chapters[0].segments if s.source.startswith("1. The myth"))
            self.assertIsNone(toc.meta.get("list_num_id"))
            self.assertIsNone(toc.meta.get("list_fmt"))

    def test_list_number_meta_and_export(self):
        from docx.oxml.ns import qn

        with tempfile.TemporaryDirectory() as directory:
            path = os.path.join(directory, "list.docx")
            doc = DocxDocument()
            doc.add_paragraph("One", style="List Number")
            doc.add_paragraph("Two", style="List Number")
            doc.save(path)
            book = read_docx(path, "en", "zh")
            segment = book.chapters[0].segments[0]
            self.assertEqual(segment.meta.get("list_fmt"), "decimal")
            self.assertIsInstance(segment.meta.get("list_num_id"), int)
            store = RunStore(os.path.join(directory, "state", "list"))
            store.save_manifest(
                {
                    "title": "list",
                    "fmt": "docx",
                    "source_lang": "en",
                    "target_lang": "zh",
                    "source_path": path,
                    "source_sha256": "x",
                    "chapters": [
                        {
                            "index": 0,
                            "title": book.chapters[0].title,
                            "status": STATUS_DONE,
                        }
                    ],
                }
            )
            chapter = book.chapters[0]
            for item in chapter.segments:
                item.target = item.source
            store.save_chapter(chapter)
            out_path = os.path.join(directory, "out.docx")
            _assemble_docx(store, out_path)
            result = DocxDocument(out_path)
            first = result.paragraphs[0]
            style = first.style
            style_name = style.name if style is not None and style.name is not None else ""
            self.assertTrue(style_name.startswith("List Number"))
            num_pr = first._p.pPr.find(qn("w:numPr")) if first._p.pPr is not None else None
            self.assertIsNotNone(num_pr)


class TestDocxStyles(unittest.TestCase):
    def test_proportional_fallback_maps_ranges(self):
        items = [
            {
                "id": "s0",
                "mode": "range",
                "source_start": 0,
                "source_end": 4,
                "color": "FF0000",
            }
        ]
        placements = proportional_range_placements("Red plain", "红的 普通", items)
        self.assertEqual(len(placements), 1)
        self.assertEqual(placements[0]["status"], "fallback")
        self.assertLess(placements[0]["target_start"], placements[0]["target_end"])
        self.assertEqual(placements[0]["color"], "FF0000")

    def test_merge_keeps_good_spans_and_inherits_style(self):
        items = [
            {
                "id": "s0",
                "mode": "range",
                "source_start": 0,
                "source_end": 10,
                "bold": True,
            },
            {
                "id": "s1",
                "mode": "range",
                "source_start": 20,
                "source_end": 30,
                "italic": True,
            },
        ]
        placements = [
            {
                "id": "s0",
                "mode": "range",
                "target_start": 0,
                "target_end": 4,
                "status": "aligned",
                "method": "llm_markers",
            },
            {
                "id": "s1",
                "mode": "range",
                "target_start": 10,
                "target_end": 10,
                "status": "fallback",
                "method": "paragraph_end",
            },
        ]
        merged, any_fallback = merge_align_results("x" * 40, "y" * 20, items, placements)
        self.assertTrue(any_fallback)
        self.assertEqual(merged[0]["method"], "llm_markers")
        self.assertEqual(merged[0]["bold"], True)
        self.assertEqual(merged[1]["method"], "proportional_source_range")
        self.assertEqual(merged[1]["italic"], True)

    def test_assemble_applies_uniform_bold(self):
        with tempfile.TemporaryDirectory() as directory:
            path = os.path.join(directory, "bold.docx")
            src = DocxDocument()
            paragraph = src.add_paragraph()
            run = paragraph.add_run("Hello")
            run.bold = True
            run.font.name = "Times New Roman"
            src.save(path)
            book = read_docx(path, "en", "zh")
            store = RunStore(os.path.join(directory, "state", "bold"))
            store.save_manifest(
                {
                    "title": "bold",
                    "fmt": "docx",
                    "source_lang": "en",
                    "target_lang": "zh",
                    "source_path": path,
                    "source_sha256": "x",
                    "chapters": [
                        {
                            "index": 0,
                            "title": book.chapters[0].title,
                            "status": STATUS_DONE,
                        }
                    ],
                }
            )
            chapter = book.chapters[0]
            chapter.segments[0].target = "你好"
            store.save_chapter(chapter)
            out_path = os.path.join(directory, "out.docx")
            _assemble_docx(store, out_path)
            result = DocxDocument(out_path)
            runs = [run for p in result.paragraphs for run in p.runs if run.text.strip()]
            self.assertTrue(runs)
            self.assertTrue(runs[0].bold)
            self.assertEqual(runs[0].font.name, "宋体")

    def test_untranslated_source_fallback_skips_song_font(self):
        with tempfile.TemporaryDirectory() as directory:
            path = os.path.join(directory, "src.docx")
            src = DocxDocument()
            paragraph = src.add_paragraph()
            run = paragraph.add_run("Hello")
            run.font.name = "Times New Roman"
            src.save(path)
            book = read_docx(path, "en", "zh")
            store = RunStore(os.path.join(directory, "state", "src"))
            store.save_manifest(
                {
                    "title": "src",
                    "fmt": "docx",
                    "source_lang": "en",
                    "target_lang": "zh",
                    "source_path": path,
                    "source_sha256": "x",
                    "chapters": [
                        {
                            "index": 0,
                            "title": book.chapters[0].title,
                            "status": STATUS_DONE,
                        }
                    ],
                }
            )
            chapter = book.chapters[0]
            # 未翻译：target 为空，写出回退原文
            chapter.segments[0].target = None
            store.save_chapter(chapter)
            out_path = os.path.join(directory, "out.docx")
            _assemble_docx(store, out_path)
            result = DocxDocument(out_path)
            runs = [run for p in result.paragraphs for run in p.runs if run.text.strip()]
            self.assertTrue(runs)
            self.assertEqual(runs[0].text, "Hello")
            self.assertNotEqual(runs[0].font.name, "宋体")

    def test_assemble_preserves_center_and_mixed_bold_without_placements(self):
        """未跑样式对齐时，导出仍应用居中与混排加粗（比例回退）。"""
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import RGBColor

        with tempfile.TemporaryDirectory() as directory:
            path = os.path.join(directory, "mixed.docx")
            src = DocxDocument()
            title = src.add_heading("Title", level=1)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph = src.add_paragraph()
            bold = paragraph.add_run("Adam Kuper")
            bold.bold = True
            bold.font.color.rgb = RGBColor(0x11, 0x22, 0x33)
            paragraph.add_run(" is a writer.")
            src.save(path)
            book = read_docx(path, "en", "zh")
            self.assertEqual(book.chapters[0].segments[0].meta.get("align"), "center")
            store = RunStore(os.path.join(directory, "state", "mixed"))
            store.save_manifest(
                {
                    "title": "mixed",
                    "fmt": "docx",
                    "source_lang": "en",
                    "target_lang": "zh",
                    "source_path": path,
                    "source_sha256": "x",
                    "chapters": [
                        {
                            "index": 0,
                            "title": book.chapters[0].title,
                            "status": STATUS_DONE,
                        }
                    ],
                }
            )
            chapter = book.chapters[0]
            for segment in chapter.segments:
                segment.target = segment.source
            store.save_chapter(chapter)
            out_path = os.path.join(directory, "out.docx")
            _assemble_docx(store, out_path)
            result = DocxDocument(out_path)
            self.assertEqual(result.paragraphs[0].alignment, WD_ALIGN_PARAGRAPH.CENTER)
            adam = next(p for p in result.paragraphs if p.text.startswith("Adam Kuper"))
            first = next(run for run in adam.runs if run.text.startswith("Adam"))
            self.assertTrue(first.bold)
            self.assertEqual(str(first.font.color.rgb), "112233")


class TestDocxAssemble(unittest.TestCase):
    def test_assemble_rebuilds_headings_and_table(self):
        with tempfile.TemporaryDirectory() as directory:
            path = os.path.join(directory, "sample.docx")
            _write_sample_docx(path)
            book = read_docx(path, "en", "zh")
            store = RunStore(os.path.join(directory, "state", "sample"))
            store.save_manifest(
                {
                    "title": "sample",
                    "fmt": "docx",
                    "source_lang": "en",
                    "target_lang": "zh",
                    "source_path": path,
                    "source_sha256": "x",
                    "chapters": [
                        {
                            "index": 0,
                            "title": book.chapters[0].title,
                            "status": STATUS_DONE,
                            "title_translated": "第一章",
                        }
                    ],
                }
            )
            chapter = book.chapters[0]
            for segment in chapter.segments:
                segment.target = f"译:{segment.source}"
            store.save_chapter(chapter)
            out_path = os.path.join(directory, "out.docx")
            written = _assemble_docx(store, out_path)
            self.assertEqual(written, out_path)
            result = DocxDocument(out_path)
            texts = [p.text for p in result.paragraphs if p.text.strip()]
            self.assertIn("译:Chapter One", texts)
            self.assertIn("译:Hello world.", texts)
            self.assertEqual(len(result.tables), 1)
            self.assertEqual(result.tables[0].cell(0, 0).text.strip(), "译:A")
            self.assertEqual(result.tables[0].cell(1, 1).text.strip(), "译:D")

            via_writer = assemble(store, path, out_format="docx")
            self.assertTrue(via_writer.endswith(".docx"))
            self.assertTrue(os.path.isfile(via_writer))


class TestDocxCliDefaults(unittest.TestCase):
    def test_resolve_output_format_defaults(self):
        self.assertEqual(_resolve_output_format("a.docx", None), "docx")
        self.assertEqual(_resolve_output_format("a.epub", None), "epub")
        self.assertEqual(_resolve_output_format("a.docx", "epub"), "epub")
        self.assertEqual(_resolve_output_format("a.txt", "docx"), "docx")

    def test_translate_docx_defaults_out_format(self):
        captured: dict = {}

        class FakeOrchestrator:
            def __init__(self, config, client=None):
                del client
                captured["config"] = config

            def run_all(self, input_path, **kwargs):
                captured["input"] = input_path
                captured["kwargs"] = kwargs
                store = type(
                    "S",
                    (),
                    {"run_dir": "state/sample", "load_usage": staticmethod(lambda: {})},
                )()
                return {
                    "store": store,
                    "outputs": ["output/sample.zh.docx"],
                    "output": "output/sample.zh.docx",
                    "report": {
                        "summary": {
                            "chapters_done": 1,
                            "chapters_total": 1,
                            "terms": 0,
                        }
                    },
                }

        with tempfile.TemporaryDirectory() as directory:
            path = os.path.join(directory, "sample.docx")
            _write_sample_docx(path)
            with (
                patch(
                    "trans_novel.cli._load_config",
                    return_value=Config.from_dict({"llm": {"provider": "fake"}}),
                ),
                patch("trans_novel.pipeline.orchestrator.Orchestrator", FakeOrchestrator),
            ):
                result = CliRunner().invoke(app, ["translate", path])
        self.assertEqual(result.exit_code, 0, result.output)
        self.assertEqual(captured["kwargs"].get("out_format"), "docx")

    def test_format_docx_is_accepted(self):
        cfg = Config.from_dict({"llm": {"provider": "fake"}})
        captured: dict = {}

        class FakeOrchestrator:
            def __init__(self, config, client=None):
                del client

            def run_all(self, input_path, **kwargs):
                captured["kwargs"] = kwargs
                store = type(
                    "S",
                    (),
                    {"run_dir": "state/sample", "load_usage": staticmethod(lambda: {})},
                )()
                return {
                    "store": store,
                    "outputs": ["output/book.zh.docx"],
                    "output": "output/book.zh.docx",
                    "report": {
                        "summary": {
                            "chapters_done": 1,
                            "chapters_total": 1,
                            "terms": 0,
                        }
                    },
                }

        with tempfile.TemporaryDirectory() as directory:
            path = os.path.join(directory, "book.txt")
            with open(path, "w", encoding="utf-8") as handle:
                handle.write("Hello.\n")
            with (
                patch("trans_novel.cli._load_config", return_value=cfg),
                patch("trans_novel.pipeline.orchestrator.Orchestrator", FakeOrchestrator),
            ):
                result = CliRunner().invoke(app, ["translate", path, "--format", "docx"])
        self.assertEqual(result.exit_code, 0, result.output)
        self.assertEqual(captured["kwargs"].get("out_format"), "docx")


if __name__ == "__main__":
    unittest.main()
