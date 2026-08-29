"""从 RunStore 重建 Word .docx：标题导航 + 段落样式 + 简易表格。

- ``meta.docx_style``：整段同质，直接套到译文 run（不经 AI）
- ``meta.docx_styles.placements``：混排对齐结果，按 target 偏移切 run
"""

from __future__ import annotations

from typing import Any

from docx import Document as open_docx
from docx.document import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from ..ingest.docx_reader import _text_has_visible_list_prefix
from ..ingest.models import KIND_HEADING, Chapter
from ..pipeline.docx_styles import proportional_range_placements
from ..pipeline.runstore import RunStore
from .writer_common import (
    _bilingual_source,
    _ch_title,
    _manifest_target_lang,
    _seg_text,
)

# 译文侧不沿用原文西文字体；中文目标默认宋体（含东亚字形）。
_ZH_FONT = "宋体"

_ALIGN_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "both": WD_ALIGN_PARAGRAPH.JUSTIFY,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    "distribute": WD_ALIGN_PARAGRAPH.DISTRIBUTE,
}


def _set_outline_level(paragraph, level: int) -> None:
    """确保段落带 outlineLvl，便于 Word 导航窗格。"""
    level = max(1, min(9, level))
    p_pr = paragraph._p.get_or_add_pPr()  # noqa: SLF001
    outline = p_pr.find(qn("w:outlineLvl"))
    if outline is None:
        outline = p_pr.makeelement(qn("w:outlineLvl"), {})
        p_pr.append(outline)
    outline.set(qn("w:val"), str(level - 1))


def _set_run_font(run, font_name: str) -> None:
    """设置 run 的 ascii/hAnsi/eastAsia 字体，避免只改西文名。"""
    name = font_name.strip()
    if not name:
        return
    run.font.name = name
    r_pr = run._r.get_or_add_rPr()  # noqa: SLF001
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        r_fonts.set(qn(attr), name)


def _target_output_font(target_lang: str | None) -> str | None:
    """目标语言对应的写出字体；中文用宋体，其它语言不强行改字体。"""
    normalized = (target_lang or "zh").strip().lower().replace("_", "-")
    if normalized == "zh" or normalized.startswith("zh-"):
        return _ZH_FONT
    return None


def _font_for_text(
    output_font: str | None,
    *,
    source: str,
    output_text: str,
    is_source_side: bool = False,
) -> str | None:
    """译文用目标字体；未翻译回退原文或双语原文侧不套宋体。"""
    if is_source_side or not output_font:
        return None
    if not output_text.strip() or output_text == source:
        return None
    return output_font


def _apply_run_style(
    run,
    style: dict[str, Any] | None,
    *,
    output_font: str | None = None,
) -> None:
    """把 meta 中的字符样式应用到 run；不沿用原文 font，改用 output_font。"""
    if style:
        if "bold" in style:
            run.bold = bool(style["bold"])
        if "italic" in style:
            run.italic = bool(style["italic"])
        if style.get("underline"):
            run.underline = True
        size_pt = style.get("size_pt")
        if isinstance(size_pt, (int, float)) and size_pt > 0:
            run.font.size = Pt(float(size_pt))
        color = style.get("color")
        if isinstance(color, str) and len(color) >= 6:
            try:
                run.font.color.rgb = RGBColor.from_string(color[-6:])
            except (ValueError, TypeError):
                pass
    if output_font:
        _set_run_font(run, output_font)


def _set_run_color_value(run, value: str) -> None:
    """写入直写颜色，清掉 themeColor，避免 Heading 默认主题蓝。"""
    r_pr = run._r.get_or_add_rPr()  # noqa: SLF001
    for child in list(r_pr):
        if child.tag == qn("w:color"):
            r_pr.remove(child)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), value)
    r_pr.append(color)


def _neutralize_heading_theme_color(
    paragraph,
    style: dict[str, Any] | None,
) -> None:
    """python-docx 默认 Heading 带 accent 蓝；无显式颜色时改为黑色。"""
    explicit = None
    if isinstance(style, dict):
        color = style.get("color")
        if isinstance(color, str) and len(color) >= 6:
            explicit = color[-6:]
    value = explicit or "000000"
    for run in paragraph.runs:
        if run.text:
            _set_run_color_value(run, value)


def _list_style_name(list_fmt: str | None, ilvl: int) -> str:
    """映射到 python-docx 内置列表样式名。"""
    level = max(0, min(2, int(ilvl)))
    bullet = (list_fmt or "").lower() in {"bullet", "none"}
    if bullet:
        return "List Bullet" if level == 0 else f"List Bullet {level + 1}"
    return "List Number" if level == 0 else f"List Number {level + 1}"


def _abstract_num_id_for_style(doc: DocxDocument, style_name: str) -> int | None:
    """从段落样式上的 numPr 找到 abstractNumId。"""
    try:
        style = doc.styles[style_name]
    except KeyError:
        return None
    try:
        num_pr = style.element.pPr.numPr  # noqa: SLF001
        num_id = int(num_pr.numId.val)
    except (AttributeError, TypeError, ValueError):
        return None
    try:
        numbering = doc.part.numbering_part._element  # noqa: SLF001
    except (AttributeError, ValueError, KeyError):
        return None
    for num in numbering.findall(qn("w:num")):
        if num.get(qn("w:numId")) == str(num_id):
            abs_el = num.find(qn("w:abstractNumId"))
            if abs_el is not None:
                try:
                    return int(abs_el.get(qn("w:val")))
                except (TypeError, ValueError):
                    return None
    return None


def _next_num_id(numbering_root) -> int:
    used = []
    for num in numbering_root.findall(qn("w:num")):
        raw = num.get(qn("w:numId"))
        if raw is not None:
            try:
                used.append(int(raw))
            except ValueError:
                continue
    return (max(used) + 1) if used else 1


def _restart_list_numbering(doc: DocxDocument, paragraph, *, style_name: str, ilvl: int) -> None:
    """为列表首项新建带 startOverride 的 numId，使各组列表从 1 重开。"""
    abstract_id = _abstract_num_id_for_style(doc, style_name)
    if abstract_id is None:
        return
    try:
        numbering = doc.part.numbering_part._element  # noqa: SLF001
    except (AttributeError, ValueError, KeyError):
        return
    new_id = _next_num_id(numbering)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(new_id))
    abs_ref = OxmlElement("w:abstractNumId")
    abs_ref.set(qn("w:val"), str(abstract_id))
    num.append(abs_ref)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), str(max(0, ilvl)))
    start = OxmlElement("w:startOverride")
    start.set(qn("w:val"), "1")
    override.append(start)
    num.append(override)
    numbering.append(num)

    p_pr = paragraph._p.get_or_add_pPr()  # noqa: SLF001
    num_pr = p_pr.get_or_add_numPr()
    num_pr.get_or_add_ilvl().val = max(0, ilvl)
    num_pr.get_or_add_numId().val = new_id


def _style_slices(
    text: str,
    style: dict[str, Any] | None,
    placements: list[dict[str, Any]] | None,
) -> list[tuple[str, dict[str, Any] | None]]:
    """把文本切成 (fragment, style) 列表；无混排时整段一个切片。"""
    if not text:
        return []
    if not placements:
        return [(text, style)]
    bounds = [0, len(text)]
    usable: list[dict[str, Any]] = []
    for row in placements:
        start = row.get("target_start")
        end = row.get("target_end")
        if not isinstance(start, int) or not isinstance(end, int):
            continue
        start = max(0, min(len(text), start))
        end = max(start, min(len(text), end))
        if start >= end:
            continue
        usable.append({**row, "target_start": start, "target_end": end})
        bounds.extend((start, end))
    if not usable:
        return [(text, style)]
    cuts = sorted(set(bounds))
    slices: list[tuple[str, dict[str, Any] | None]] = []
    for left, right in zip(cuts, cuts[1:]):
        if left >= right:
            continue
        fragment = text[left:right]
        matched: dict[str, Any] | None = None
        for row in usable:
            if row["target_start"] <= left and right <= row["target_end"]:
                matched = {
                    key: row[key]
                    for key in ("bold", "italic", "underline", "color", "size_pt", "font")
                    if key in row
                }
                break
        slices.append((fragment, matched or style))
    return slices or [(text, style)]


def _apply_paragraph_align(paragraph, align: str | None) -> None:
    """应用段落对齐。"""
    if not align:
        return
    value = _ALIGN_MAP.get(str(align).strip().lower())
    if value is not None:
        paragraph.alignment = value


def _apply_paragraph_shade(paragraph, shade: str | None) -> None:
    """应用段落底纹填充色。"""
    if not shade:
        return
    p_pr = paragraph._p.get_or_add_pPr()  # noqa: SLF001
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = p_pr.makeelement(qn("w:shd"), {})
        p_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), str(shade).upper()[-6:])


def _fill_paragraph(
    paragraph,
    text: str,
    *,
    style: dict[str, Any] | None = None,
    placements: list[dict[str, Any]] | None = None,
    align: str | None = None,
    shade: str | None = None,
    output_font: str | None = None,
    dim: bool = False,
) -> None:
    """清空并按样式切片写入段落。"""
    paragraph.clear()
    _apply_paragraph_align(paragraph, align)
    _apply_paragraph_shade(paragraph, shade)
    for fragment, frag_style in _style_slices(text, style, placements):
        run = paragraph.add_run(fragment)
        _apply_run_style(run, frag_style, output_font=output_font)
        if dim:
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            try:
                run.font.highlight_color = WD_COLOR_INDEX.GRAY_25
            except (AttributeError, ValueError):
                pass


def _add_heading(
    doc: DocxDocument,
    text: str,
    level: int,
    *,
    style: dict[str, Any] | None = None,
    placements: list[dict[str, Any]] | None = None,
    align: str | None = None,
    shade: str | None = None,
    output_font: str | None = None,
) -> None:
    level = max(1, min(9, level))
    style_name = f"Heading {level}"
    try:
        paragraph = doc.add_heading("", level=level)
    except (KeyError, ValueError):
        paragraph = doc.add_paragraph("")
        try:
            paragraph.style = style_name
        except (KeyError, ValueError):
            pass
    _set_outline_level(paragraph, level)
    _fill_paragraph(
        paragraph,
        text,
        style=style,
        placements=placements,
        align=align,
        shade=shade,
        output_font=output_font,
    )
    # 覆盖模板 Heading 的 accent 主题蓝：原文无显式色则用黑色
    _neutralize_heading_theme_color(paragraph, style)


def _add_normal(
    doc: DocxDocument,
    text: str,
    *,
    style: dict[str, Any] | None = None,
    placements: list[dict[str, Any]] | None = None,
    align: str | None = None,
    shade: str | None = None,
    output_font: str | None = None,
    dim: bool = False,
    list_fmt: str | None = None,
    list_ilvl: int = 0,
    list_restart: bool = False,
) -> None:
    if list_fmt is not None:
        style_name = _list_style_name(list_fmt, list_ilvl)
        try:
            paragraph = doc.add_paragraph(style=style_name)
        except KeyError:
            paragraph = doc.add_paragraph()
            style_name = "List Number"
        _fill_paragraph(
            paragraph,
            text,
            style=style,
            placements=placements,
            align=align,
            shade=shade,
            output_font=output_font,
            dim=dim,
        )
        if list_restart:
            _restart_list_numbering(doc, paragraph, style_name=style_name, ilvl=list_ilvl)
        return

    paragraph = doc.add_paragraph()
    _fill_paragraph(
        paragraph,
        text,
        style=style,
        placements=placements,
        align=align,
        shade=shade,
        output_font=output_font,
        dim=dim,
    )


def _add_bilingual_paragraphs(
    doc: DocxDocument,
    source: str,
    target: str,
    order: str,
    *,
    style: dict[str, Any] | None = None,
    placements: list[dict[str, Any]] | None = None,
    align: str | None = None,
    shade: str | None = None,
    output_font: str | None = None,
) -> None:
    src = _bilingual_source(source, target)
    target_font = _font_for_text(output_font, source=source, output_text=target)
    if not src:
        _add_normal(
            doc,
            target,
            style=style,
            placements=placements,
            align=align,
            shade=shade,
            output_font=target_font,
        )
        return
    # 译文侧可用宋体；原文侧保持默认/不强制目标字体
    if order == "source_first":
        _add_normal(doc, src, dim=False, align=align, shade=shade, output_font=None)
        _add_normal(
            doc,
            target,
            style=style,
            placements=placements,
            align=align,
            shade=shade,
            output_font=target_font,
        )
    else:
        _add_normal(
            doc,
            target,
            style=style,
            placements=placements,
            align=align,
            shade=shade,
            output_font=target_font,
        )
        _add_normal(doc, src, dim=True, align=align, shade=shade, output_font=None)


def _segment_style_payload(
    meta: dict[str, Any],
    *,
    source: str,
    output_text: str,
) -> tuple[dict[str, Any] | None, list | None, str | None, str | None]:
    """返回 (整段样式, 混排 placements, align, shade)。

    若混排尚无 placements（未对齐或未翻译就导出），按源文偏移比例映到当前写出文本，
    这样「Adam Kuper」加粗等在回退原文导出时也不会丢。
    """
    align = meta.get("align") if isinstance(meta.get("align"), str) else None
    shade = meta.get("shade") if isinstance(meta.get("shade"), str) else None
    uniform = meta.get("docx_style")
    if isinstance(uniform, dict) and uniform:
        return uniform, None, align, shade
    styles = meta.get("docx_styles")
    if not isinstance(styles, dict):
        return None, None, align, shade
    placements = styles.get("placements")
    if isinstance(placements, list) and placements:
        return None, placements, align, shade
    items = styles.get("items")
    if isinstance(items, list) and items:
        return None, proportional_range_placements(source, output_text, items), align, shade
    return None, None, align, shade


def _flush_table(
    doc: DocxDocument,
    cells: dict[tuple[int, int], tuple[str, str, dict[str, Any]]],
    rows: int,
    cols: int,
    *,
    bilingual: bool,
    order: str,
    output_font: str | None = None,
) -> None:
    table = doc.add_table(rows=rows, cols=cols)
    try:
        table.style = "Table Grid"
    except (KeyError, ValueError):
        pass
    for r in range(rows):
        for c in range(cols):
            target, source, meta = cells.get((r, c), ("", "", {}))
            style, placements, align, shade = _segment_style_payload(
                meta, source=source, output_text=target
            )
            target_font = _font_for_text(output_font, source=source, output_text=target)
            cell = table.cell(r, c)
            cell.text = ""
            paragraph = cell.paragraphs[0]
            if bilingual:
                src = _bilingual_source(source, target)
                if src:
                    if order == "source_first":
                        _fill_paragraph(
                            paragraph,
                            src,
                            align=align,
                            shade=shade,
                            output_font=None,
                        )
                        paragraph.add_run("\n")
                        # 译文另起逻辑：简化为同一段内第二行
                        for fragment, frag_style in _style_slices(target, style, placements):
                            run = paragraph.add_run(fragment)
                            _apply_run_style(run, frag_style, output_font=target_font)
                    else:
                        _fill_paragraph(
                            paragraph,
                            target,
                            style=style,
                            placements=placements,
                            align=align,
                            shade=shade,
                            output_font=target_font,
                        )
                        paragraph.add_run("\n")
                        dim_run = paragraph.add_run(src)
                        dim_run.font.size = Pt(9)
                        dim_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
                else:
                    _fill_paragraph(
                        paragraph,
                        target,
                        style=style,
                        placements=placements,
                        align=align,
                        shade=shade,
                        output_font=target_font,
                    )
            else:
                _fill_paragraph(
                    paragraph,
                    target,
                    style=style,
                    placements=placements,
                    align=align,
                    shade=shade,
                    output_font=target_font,
                )


def _emit_chapter_blocks(
    doc: DocxDocument,
    chapter: Chapter,
    *,
    bilingual: bool,
    order: str,
    output_font: str | None = None,
) -> None:
    """按段顺序写出；连续同 table_id 聚合成一张表；cont 续段并回上一段。"""
    i = 0
    segs = chapter.segments
    last_list_num_id: int | None = None
    while i < len(segs):
        seg = segs[i]
        meta = seg.meta if isinstance(seg.meta, dict) else {}
        table_id = meta.get("table_id")
        if isinstance(table_id, int):
            cells: dict[tuple[int, int], tuple[str, str, dict[str, Any]]] = {}
            rows = int(meta.get("rows") or 1)
            cols = int(meta.get("cols") or 1)
            while i < len(segs):
                cur = segs[i]
                cur_meta = cur.meta if isinstance(cur.meta, dict) else {}
                if cur_meta.get("table_id") != table_id:
                    break
                r = int(cur_meta.get("row") or 0)
                c = int(cur_meta.get("col") or 0)
                rows = max(rows, int(cur_meta.get("rows") or rows))
                cols = max(cols, int(cur_meta.get("cols") or cols))
                cells[(r, c)] = (_seg_text(cur), cur.source, cur_meta)
                i += 1
            _flush_table(
                doc,
                cells,
                rows,
                cols,
                bilingual=bilingual,
                order=order,
                output_font=output_font,
            )
            last_list_num_id = None
            continue

        if not seg.source.strip() and not (seg.target and seg.target.strip()):
            i += 1
            continue

        target_parts = [_seg_text(seg)]
        source_parts = [seg.source]
        kind = seg.kind
        heading_level = 1
        if kind == KIND_HEADING:
            raw_level = meta.get("heading_level", 1)
            heading_level = raw_level if isinstance(raw_level, int) else 1
        style_meta = meta
        i += 1
        while i < len(segs):
            nxt = segs[i]
            nxt_meta = nxt.meta if isinstance(nxt.meta, dict) else {}
            if not nxt.cont or nxt_meta.get("table_id") is not None:
                break
            target_parts.append(_seg_text(nxt))
            source_parts.append(nxt.source)
            i += 1
        target = "".join(target_parts)
        source = "".join(source_parts)
        style, placements, align, shade = _segment_style_payload(
            style_meta, source=source, output_text=target
        )
        text_font = _font_for_text(output_font, source=source, output_text=target)
        list_num_id = style_meta.get("list_num_id")
        list_ilvl = style_meta.get("list_ilvl")
        list_fmt = style_meta.get("list_fmt")
        is_list = (
            isinstance(list_num_id, int)
            and list_num_id > 0
            and isinstance(list_fmt, str)
            and not _text_has_visible_list_prefix(target)
            and not _text_has_visible_list_prefix(source)
        )
        if kind == KIND_HEADING:
            _add_heading(
                doc,
                target,
                heading_level,
                style=style,
                placements=placements,
                align=align,
                shade=shade,
                output_font=text_font,
            )
            last_list_num_id = None
        elif bilingual:
            _add_bilingual_paragraphs(
                doc,
                source,
                target,
                order,
                style=style,
                placements=placements,
                align=align,
                shade=shade,
                output_font=output_font,
            )
            last_list_num_id = None
        elif is_list:
            restart = list_num_id != last_list_num_id
            _add_normal(
                doc,
                target,
                style=style,
                placements=placements,
                align=align,
                shade=shade,
                output_font=text_font,
                list_fmt=list_fmt,
                list_ilvl=int(list_ilvl) if isinstance(list_ilvl, int) else 0,
                list_restart=restart,
            )
            last_list_num_id = list_num_id
        else:
            _add_normal(
                doc,
                target,
                style=style,
                placements=placements,
                align=align,
                shade=shade,
                output_font=text_font,
            )
            last_list_num_id = None


def _assemble_docx(
    store: RunStore,
    out_path: str,
    *,
    bilingual: bool = False,
    order: str = "target_first",
) -> str:
    """按章节重建 .docx；标题带 outline，样式与表格按 meta 重建。"""
    manifest = store.load_manifest()
    output_font = _target_output_font(_manifest_target_lang(manifest))
    doc = open_docx()
    if doc.paragraphs:
        p0 = doc.paragraphs[0]
        if not p0.text.strip():
            p0.clear()

    first_block = True
    for c in manifest["chapters"]:
        chapter = store.load_chapter(c["index"])
        has_h1 = any(
            s.kind == KIND_HEADING
            and isinstance(s.meta, dict)
            and int(s.meta.get("heading_level") or 1) == 1
            for s in chapter.segments
        )
        title = _ch_title(c)
        if title and not has_h1 and chapter.meta.get("explicit_title"):
            if first_block and doc.paragraphs and not doc.paragraphs[0].text:
                pass
            _add_heading(doc, title, 1, output_font=output_font)
        _emit_chapter_blocks(
            doc,
            chapter,
            bilingual=bilingual,
            order=order,
            output_font=output_font,
        )
        first_block = False

    body = doc.element.body
    for child in list(body):
        if child.tag == qn("w:p") and not (child.text or "").strip():
            texts = [node.text for node in child.iter(qn("w:t")) if node.text]
            if not any(texts) and body.index(child) == 0:
                body.remove(child)
            break

    doc.save(out_path)
    return out_path
